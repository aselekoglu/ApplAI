from __future__ import annotations

import os
import re
import json
from docx import Document


def _load_template_config(template_config):
    if not template_config:
        return {}
    if isinstance(template_config, dict):
        return template_config
    if isinstance(template_config, str) and os.path.exists(template_config):
        try:
            with open(template_config, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as exc:
            print(f"[Renderer] Failed to read template config ({exc}).")
            return {}
    return {}


def _resolve_headers(config: dict, key: str, defaults: list[str]) -> list[str]:
    headers = config.get(key, defaults)
    if isinstance(headers, str):
        return [headers]
    if isinstance(headers, list):
        return [h for h in headers if isinstance(h, str) and h.strip()]
    return defaults


def generate_tailored_document(
    template_path: str,
    output_path: str,
    tailored_data: dict,
    template_config=None,
    max_pages: int = 2,
):
    """
    Injects tailored CV content into a DOCX template.

    Accepts two formats for tailored_data:

    1. New format (from WorkflowResult.tailored_output):
       - 'tailored_output': TailoredOutput Pydantic object

    2. Legacy format (dict with string lists):
       - 'profile_bullets': list[str]
       - 'experience_highlights': list[str]
       - 'skills_to_highlight': list[str]  (optional)
    """
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return False

    merged_config = _load_template_config(template_config or tailored_data.get("template_config"))
    if isinstance(tailored_data.get("template_config"), dict):
        merged_config = {**merged_config, **tailored_data.get("template_config", {})}

    doc = Document(template_path)

    #  Resolve bullet lists from either format 
    profile_bullets: list[str] = []
    experience_highlights: list[str] = []
    education_highlights: list[str] = []
    project_highlights: list[str] = []
    skills_to_highlight: list[str] = []

    tailored_obj = tailored_data.get("tailored_output")
    if tailored_obj is not None:
        # New structured format: TailoredOutput Pydantic model
        profile_bullets = [
            s.new_text or s.original_text
            for s in tailored_obj.profile_selections
            if s.action != "deselect" and (s.new_text or s.original_text)
        ]
        experience_highlights = [
            s.new_text or s.original_text
            for s in tailored_obj.experience_selections
            if s.action != "deselect" and (s.new_text or s.original_text)
        ]
        education_highlights = [
            s.new_text or s.original_text
            for s in getattr(tailored_obj, "education_selections", [])
            if s.action != "deselect" and (s.new_text or s.original_text)
        ]
        project_highlights = [
            s.new_text or s.original_text
            for s in getattr(tailored_obj, "project_selections", [])
            if s.action != "deselect" and (s.new_text or s.original_text)
        ]
        skills_to_highlight = tailored_obj.skills_to_highlight or []
    else:
        # Legacy format
        profile_bullets = tailored_data.get("profile_bullets", [])
        experience_highlights = tailored_data.get("experience_highlights", [])
        education_highlights = tailored_data.get("education_highlights", [])
        project_highlights = tailored_data.get("project_highlights", [])
        skills_to_highlight = tailored_data.get("skills_to_highlight", [])

        tailored_raw = tailored_data.get("tailored_raw", "")
        if not profile_bullets and tailored_raw:
            extracted = _extract_section(tailored_raw, ["profile", "summary", "objective"])
            if extracted:
                lines = [l.strip().lstrip("-•*").strip() for l in extracted.splitlines() if l.strip()]
                profile_bullets = [l for l in lines if len(l) > 5]
        if not experience_highlights and tailored_raw:
            extracted = _extract_section(tailored_raw, ["experience", "relevant experience", "work experience"])
            if extracted:
                lines = [l.strip().lstrip("-•*").strip() for l in extracted.splitlines() if l.strip()]
                experience_highlights = [l for l in lines if len(l) > 5]

    #  Inject sections 
    profile_headers = _resolve_headers(merged_config, "profile_headers", ["PROFILE"])
    experience_headers = _resolve_headers(merged_config, "experience_headers", ["RELEVANT EXPERIENCE"])
    education_headers = _resolve_headers(merged_config, "education_headers", ["EDUCATION"])
    project_headers = _resolve_headers(merged_config, "project_headers", ["PROJECTS", "PROJECT EXPERIENCE"])
    skill_headers = _resolve_headers(merged_config, "skills_headers", ["SUMMARY OF QUALIFICATIONS", "KEY SKILLS", "SKILLS", "TECHNICAL SKILLS"])

    if profile_bullets:
        _replace_first_matching_section(doc, profile_headers, profile_bullets)

    if experience_highlights:
        _replace_sections_in_order(doc, experience_headers, experience_highlights)

    if education_highlights:
        _replace_first_matching_section(doc, education_headers, education_highlights)

    if project_highlights:
        _replace_first_matching_section(doc, project_headers, project_highlights)

    # Skills: try common section header variants
    if skills_to_highlight:
        _replace_first_matching_section(doc, skill_headers, skills_to_highlight)

    _apply_layout_guards(doc, profile_headers + experience_headers + education_headers + project_headers + skill_headers)

    # Overflow check
    total_paras = sum(1 for p in doc.paragraphs if p.text.strip())
    max_paragraphs = 80 if max_pages <= 2 else 120
    if total_paras > max_paragraphs:
        print(f"[Renderer]  Overflow risk: {total_paras} non-empty paragraphs (target <= {max_paragraphs}).")

    try:
        doc.save(output_path)
        print(f"[Renderer]  Tailored CV saved: {output_path}")
        return True
    except Exception as e:
        print(f"[Renderer]  Error saving document: {e}")
        return False


def _section_exists(doc, header: str) -> bool:
    """Check if a section header exists in the document."""
    for para in doc.paragraphs:
        if header.upper() in para.text.upper() and len(para.text.strip()) < 80:
            return True
    return False


def _extract_section(text: str, keywords: list) -> str:
    """Extract a named section from agent raw text output."""
    lines = text.splitlines()
    in_section = False
    section_lines = []
    for line in lines:
        line_lower = line.strip().lower()
        if any(kw in line_lower for kw in keywords) and len(line.strip()) < 60:
            in_section = True
            continue
        if in_section:
            if (line.strip().isupper() and len(line.strip()) > 3) or line.strip().startswith("##"):
                break
            section_lines.append(line)
    return "\n".join(section_lines).strip()


def _replace_section_bullets(doc, section_header: str, new_bullets: list):
    """
    Find a section header in the doc and replace the following
    bullet paragraphs with new_bullets, preserving run formatting.
    """
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        if section_header.upper() in para.text.upper() and len(para.text.strip()) < 80:
            header_idx = i
            break

    if header_idx is None:
        print(f"[Renderer] Section '{section_header}' not found in document.")
        return False

    # Collect bullet paragraphs after the header.
    # In this DOCX all paragraphs use style='normal'. The ONLY reliable
    # discriminator is left_indent:
    #   indent = 0      -> job title row, company name, date line  (SKIP)
    #   indent > 0      -> actual indented bullet point            (REPLACE)
    # Fallback: also accept paragraphs with an explicit bullet prefix character.
    BULLET_CHARS = {'\u2022', '\u25cf', '\u25e6', '\u25aa', '-', '*'}
    bullet_indices = []

    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # Stop at the next all-caps section header (EDUCATION, SKILLS, etc.)
        if text and text.upper() == text and len(text) < 80 and len(text) > 3 and i > header_idx + 1:
            break
        if not text:
            continue

        indent = para.paragraph_format.left_indent or 0
        has_bullet_prefix = bool(text) and text[0] in BULLET_CHARS

        # Only replace indented paragraphs or explicit bullet-prefixed lines
        if indent > 0 or has_bullet_prefix:
            bullet_indices.append(i)
        # else: job title / company / date line -- leave untouched

        if len(bullet_indices) >= 12:
            break


    if not bullet_indices:
        print(f"[Renderer] No bullet paragraphs found under '{section_header}'.")
        return False

    print(f"[Renderer] Replacing {min(len(bullet_indices), len(new_bullets))} bullets under '{section_header}'")

    # Replace paragraph text in-place, preserving run formatting
    for doc_idx, new_text in zip(bullet_indices, new_bullets):
        para = doc.paragraphs[doc_idx]
        # Clean any leading bullet chars from the new text
        clean_text = re.sub(r'^[\-•*]\s*', '', new_text.strip())
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic
            for run in para.runs:
                run.text = ""
            first_run.text = clean_text
            if font_name:
                first_run.font.name = font_name
            if font_size:
                first_run.font.size = font_size
            first_run.bold = bold
            first_run.italic = italic
        else:
            para.text = clean_text
    return True


def _replace_first_matching_section(doc, headers: list[str], new_bullets: list):
    for header in headers:
        if _section_exists(doc, header):
            if _replace_section_bullets(doc, header, new_bullets):
                return True
    print(f"[Renderer] None of these headers were found: {headers}")
    return False


def _replace_sections_in_order(doc, headers: list[str], new_bullets: list):
    """
    Replace bullet slots across multiple matching section headers.
    This helps keep multi-role experience blocks from collapsing into one list.
    """
    valid_headers = [h for h in headers if _section_exists(doc, h)]
    if not valid_headers:
        print(f"[Renderer] None of these headers were found: {headers}")
        return False
    if not new_bullets:
        return True

    chunk_size = max(1, (len(new_bullets) + len(valid_headers) - 1) // len(valid_headers))
    replaced_any = False
    cursor = 0
    for header in valid_headers:
        chunk = new_bullets[cursor:cursor + chunk_size]
        if not chunk:
            break
        replaced = _replace_section_bullets(doc, header, chunk)
        replaced_any = replaced_any or replaced
        cursor += chunk_size
    return replaced_any


def _apply_layout_guards(doc, section_headers: list[str]):
    """
    Keep section headers attached to the next paragraph to reduce split blocks
    when the CV is close to page boundaries.
    """
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if any(h.upper() in text.upper() for h in section_headers):
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.widow_control = True
            if idx + 1 < len(doc.paragraphs):
                doc.paragraphs[idx + 1].paragraph_format.widow_control = True


if __name__ == "__main__":
    mock_data = {
        "profile_bullets": [
            "Analytical Software & Data Developer with 3+ years of experience in end-to-end automation and AI-integrated workflows.",
            "Proven expertise in Python, REST APIs, and ML pipelines with a track record of delivering data-driven solutions.",
            "Collaborates effectively in Agile, cross-functional teams to deliver production-grade applications.",
        ],
        "experience_highlights": [
            "Designed and implemented API-driven middleware enabling multiple CRM platforms to operate through a centralized logic layer.",
            "Built event-driven automation pipelines using JSON payloads and HTTP endpoints for enterprise-grade workflows.",
            "Developed and maintained automation scripts to validate and synchronize high-volume operational data across systems.",
        ],
        "skills_to_highlight": ["Python", "REST APIs", "SQL", "Machine Learning", "Agile"],
    }
    base_template = "docs/Selekoglu CV 2026 - Public Sector (Algonquin Email).docx"
    generate_tailored_document(base_template, "docs/Tailored_Output.docx", mock_data)
